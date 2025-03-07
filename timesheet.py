import streamlit as st
import pandas as pd
import numpy as np
import datetime
import calendar
import os
import plotly.express as px
import plotly.graph_objects as go
from docx import Document
from io import BytesIO
import base64
from PIL import Image
import sqlite3
import hashlib
import hmac
import secrets
import time
from datetime import datetime, timedelta
import shutil
import glob
import io
import zipfile
import json

# Set a constant secure key for password verification (don't change this after setup)
# In production, you would store this in an environment variable or secret manager
SECRET_KEY = "YOUR_SECRET_KEY_REPLACE_WITH_SOMETHING_SECURE"

# Set page config
st.set_page_config(
    page_title="Team Timesheet Tracker",
    page_icon="⏱️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Authentication functions
def generate_password_hash(password):
    """Generate a secure hash for the password"""
    return hashlib.pbkdf2_hmac(
        'sha256', 
        password.encode('utf-8'), 
        SECRET_KEY.encode('utf-8'), 
        100000
    ).hex()

def verify_password(stored_hash, provided_password):
    """Verify the provided password against the stored hash"""
    new_hash = hashlib.pbkdf2_hmac(
        'sha256', 
        provided_password.encode('utf-8'), 
        SECRET_KEY.encode('utf-8'), 
        100000
    ).hex()
    return hmac.compare_digest(new_hash, stored_hash)

def init_auth():
    """Initialize authentication settings in the database"""
    conn = sqlite3.connect('timesheet_data.db')
    c = conn.cursor()
    
    # Create auth table if it doesn't exist
    c.execute('''
    CREATE TABLE IF NOT EXISTS auth_settings (
        id INTEGER PRIMARY KEY,
        password_hash TEXT NOT NULL,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )
    ''')
    
    # Check if a password is already set
    c.execute("SELECT COUNT(*) FROM auth_settings")
    if c.fetchone()[0] == 0:
        # Set a default password (change this to your desired password)
        default_password = "timesheet2025"  # You should change this!
        password_hash = generate_password_hash(default_password)
        
        c.execute("INSERT INTO auth_settings (password_hash) VALUES (?)", (password_hash,))
        conn.commit()
    
    conn.close()

def check_auth():
    """Check if user is authenticated"""
    if 'authenticated' not in st.session_state:
        st.session_state.authenticated = False
    
    if not st.session_state.authenticated:
        auth_placeholder = st.empty()
        with auth_placeholder.container():
            st.markdown("## 🔒 Authentication Required")
            st.markdown("Please enter the password to access the Team Timesheet Tracker.")
            
            password_input = st.text_input("Password", type="password")
            login_button = st.button("Login")
            
            if login_button:
                conn = sqlite3.connect('timesheet_data.db')
                c = conn.cursor()
                c.execute("SELECT password_hash FROM auth_settings LIMIT 1")
                stored_hash = c.fetchone()[0]
                conn.close()
                
                if verify_password(stored_hash, password_input):
                    st.session_state.authenticated = True
                    auth_placeholder.empty()
                    st.rerun()
                else:
                    st.error("Incorrect password. Please try again.")
        
        # Hide the rest of the app
        st.stop()

def change_password_ui():
    """UI for changing password"""
    st.subheader("Change Application Password")
    
    current_password = st.text_input("Current Password", type="password", key="current_pwd")
    new_password = st.text_input("New Password", type="password", key="new_pwd")
    confirm_password = st.text_input("Confirm New Password", type="password", key="confirm_pwd")
    
    if st.button("Update Password"):
        if new_password != confirm_password:
            st.error("New passwords don't match.")
            return
        
        if len(new_password) < 8:
            st.error("New password must be at least 8 characters long.")
            return
            
        conn = sqlite3.connect('timesheet_data.db')
        c = conn.cursor()
        c.execute("SELECT password_hash FROM auth_settings LIMIT 1")
        stored_hash = c.fetchone()[0]
        
        if verify_password(stored_hash, current_password):
            # Update password
            new_hash = generate_password_hash(new_password)
            c.execute("UPDATE auth_settings SET password_hash = ?", (new_hash,))
            conn.commit()
            conn.close()
            st.success("Password updated successfully!")
        else:
            conn.close()
            st.error("Current password is incorrect.")

# ============= BACKUP FUNCTIONS =============

def setup_auto_backup():
    """Setup automatic backup system for the database"""
    # Create backups directory if it doesn't exist
    if not os.path.exists('backups'):
        os.makedirs('backups')
    
    # Check when the last backup was made
    conn = sqlite3.connect('timesheet_data.db')
    c = conn.cursor()
    
    # Create backup_logs table if it doesn't exist
    c.execute('''
    CREATE TABLE IF NOT EXISTS backup_logs (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        backup_date TIMESTAMP,
        backup_path TEXT,
        backup_type TEXT
    )
    ''')
    
    # Check the last backup date
    c.execute("SELECT MAX(backup_date) FROM backup_logs WHERE backup_type = 'auto'")
    last_backup = c.fetchone()[0]
    
    # Convert to datetime if not None
    if last_backup:
        last_backup = datetime.strptime(last_backup, '%Y-%m-%d %H:%M:%S')
    
    # If no backup in the last 24 hours, create one
    if not last_backup or (datetime.now() - last_backup) > timedelta(hours=24):
        # Create backup filename with timestamp
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        backup_path = f"backups/timesheet_backup_{timestamp}.db"
        
        # Close the connection to make a clean copy
        conn.close()
        
        # Copy the database file
        shutil.copy2('timesheet_data.db', backup_path)
        
        # Reopen connection
        conn = sqlite3.connect('timesheet_data.db')
        c = conn.cursor()
        
        # Log the backup
        c.execute("""
        INSERT INTO backup_logs (backup_date, backup_path, backup_type) 
        VALUES (?, ?, ?)
        """, (datetime.now().strftime('%Y-%m-%d %H:%M:%S'), backup_path, 'auto'))
        
        # Keep only the last 30 auto backups
        c.execute("SELECT backup_path FROM backup_logs WHERE backup_type = 'auto' ORDER BY backup_date DESC LIMIT 30")
        keep_backups = [row[0] for row in c.fetchall()]
        
        c.execute("SELECT backup_path FROM backup_logs WHERE backup_type = 'auto' ORDER BY backup_date DESC")
        all_backups = [row[0] for row in c.fetchall()]
        
        # Delete old backups
        for backup in all_backups:
            if backup not in keep_backups and os.path.exists(backup):
                try:
                    os.remove(backup)
                    c.execute("DELETE FROM backup_logs WHERE backup_path = ?", (backup,))
                except:
                    pass  # Skip if file is in use or can't be deleted
        
    conn.commit()
    conn.close()

def create_manual_backup():
    """Create a manual backup of the database"""
    # Create backups directory if it doesn't exist
    if not os.path.exists('backups'):
        os.makedirs('backups')
    
    # Create backup filename with timestamp
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    backup_path = f"backups/timesheet_manual_backup_{timestamp}.db"
    
    # Copy the database file
    shutil.copy2('timesheet_data.db', backup_path)
    
    # Log the backup
    conn = sqlite3.connect('timesheet_data.db')
    c = conn.cursor()
    c.execute("""
    INSERT INTO backup_logs (backup_date, backup_path, backup_type) 
    VALUES (?, ?, ?)
    """, (datetime.now().strftime('%Y-%m-%d %H:%M:%S'), backup_path, 'manual'))
    conn.commit()
    conn.close()
    
    return backup_path

def get_available_backups():
    """Get a list of all available backups"""
    conn = sqlite3.connect('timesheet_data.db')
    c = conn.cursor()
    
    # Check if table exists
    c.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='backup_logs'")
    if not c.fetchone():
        conn.close()
        return pd.DataFrame()  # Return empty dataframe if table doesn't exist
    
    df = pd.read_sql_query("""
    SELECT 
        id,
        backup_date, 
        backup_path,
        backup_type,
        CASE 
            WHEN backup_type = 'auto' THEN 'Automatic' 
            WHEN backup_type = 'manual' THEN 'Manual'
            ELSE backup_type
        END as backup_type_display
    FROM backup_logs 
    ORDER BY backup_date DESC
    """, conn)
    conn.close()
    return df

def restore_from_backup(backup_path):
    """Restore the database from a backup file"""
    # First create a backup of the current state
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    pre_restore_backup = f"backups/pre_restore_backup_{timestamp}.db"
    
    # Create backups directory if it doesn't exist
    if not os.path.exists('backups'):
        os.makedirs('backups')
    
    # Copy current database to pre-restore backup
    shutil.copy2('timesheet_data.db', pre_restore_backup)
    
    # Log the pre-restore backup
    conn = sqlite3.connect('timesheet_data.db')
    c = conn.cursor()
    c.execute("""
    INSERT INTO backup_logs (backup_date, backup_path, backup_type) 
    VALUES (?, ?, ?)
    """, (datetime.now().strftime('%Y-%m-%d %H:%M:%S'), pre_restore_backup, 'pre-restore'))
    conn.commit()
    conn.close()
    
    # Copy the backup file to the main database file
    try:
        shutil.copy2(backup_path, 'timesheet_data.db')
        return True
    except Exception as e:
        return str(e)

def create_backup_zip():
    """Create a zip file containing all backups"""
    memory_file = io.BytesIO()
    
    with zipfile.ZipFile(memory_file, 'w', zipfile.ZIP_DEFLATED) as zf:
        # Add the current database
        zf.write('timesheet_data.db', 'current_timesheet_data.db')
        
        # Add all backups
        if os.path.exists('backups'):
            for backup_file in glob.glob('backups/*.db'):
                zf.write(backup_file, os.path.basename(backup_file))
    
    memory_file.seek(0)
    return memory_file

# Initialize database
def init_db():
    conn = sqlite3.connect('timesheet_data.db')
    c = conn.cursor()
    
    # Create tables if they don't exist
    c.execute('''
    CREATE TABLE IF NOT EXISTS team_members (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        name TEXT NOT NULL UNIQUE
    )
    ''')
    
    c.execute('''
    CREATE TABLE IF NOT EXISTS timesheet_entries (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        team_member_id INTEGER,
        entry_date DATE,
        hours REAL,
        notes TEXT,
        status TEXT DEFAULT 'Pending',
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (team_member_id) REFERENCES team_members (id)
    )
    ''')
    
    # Create a flag table to track if initial data has been loaded
    c.execute('''
    CREATE TABLE IF NOT EXISTS app_settings (
        key TEXT PRIMARY KEY,
        value TEXT
    )
    ''')
    
    # Check if initial data has been loaded
    c.execute("SELECT value FROM app_settings WHERE key = 'initial_data_loaded'")
    result = c.fetchone()
    
    # Only add default data if it hasn't been loaded before
    if not result or result[0] != 'true':
        c.execute("SELECT COUNT(*) FROM team_members")
        if c.fetchone()[0] == 0:
            default_members = ["Alex Johnson", "Sarah Lee", "Michael Davis", "Emma Wilson"]
            for member in default_members:
                c.execute("INSERT INTO team_members (name) VALUES (?)", (member,))
        
        # Mark that initial data has been loaded
        c.execute("INSERT OR REPLACE INTO app_settings (key, value) VALUES (?, ?)", 
                 ('initial_data_loaded', 'true'))
    
    conn.commit()
    conn.close()

# Function to get team members from database
def get_team_members():
    conn = sqlite3.connect('timesheet_data.db')
    df = pd.read_sql_query("SELECT id, name FROM team_members ORDER BY name", conn)
    conn.close()
    return df

# Function to save timesheet entry
def save_timesheet_entry(team_member_id, entry_date, hours, notes):
    conn = sqlite3.connect('timesheet_data.db')
    c = conn.cursor()
    
    # Check if entry already exists for this combination
    c.execute("""
    SELECT id FROM timesheet_entries 
    WHERE team_member_id = ? AND entry_date = ?
    """, (team_member_id, entry_date))
    
    result = c.fetchone()
    
    if result:
        # Update existing entry
        c.execute("""
        UPDATE timesheet_entries 
        SET hours = ?, notes = ?, updated_at = CURRENT_TIMESTAMP
        WHERE id = ?
        """, (hours, notes, result[0]))
        entry_id = result[0]
    else:
        # Insert new entry
        c.execute("""
        INSERT INTO timesheet_entries 
        (team_member_id, entry_date, hours, notes) 
        VALUES (?, ?, ?, ?)
        """, (team_member_id, entry_date, hours, notes))
        entry_id = c.lastrowid
    
    conn.commit()
    conn.close()
    return entry_id

# Function to get timesheet entries for a date range
def get_timesheet_entries(start_date, end_date, team_member_id=None):
    conn = sqlite3.connect('timesheet_data.db')
    
    query = """
    SELECT 
        e.id, 
        tm.name as team_member, 
        e.entry_date, 
        e.hours, 
        e.notes, 
        e.status
    FROM timesheet_entries e
    JOIN team_members tm ON e.team_member_id = tm.id
    WHERE e.entry_date BETWEEN ? AND ?
    """
    
    params = [start_date, end_date]
    
    if team_member_id:
        query += " AND e.team_member_id = ?"
        params.append(team_member_id)
    
    query += " ORDER BY e.entry_date, tm.name"
    
    df = pd.read_sql_query(query, conn, params=params)
    conn.close()
    
    # Convert date string to datetime
    if not df.empty:
        df['entry_date'] = pd.to_datetime(df['entry_date'])
    
    return df

# Function to get entry for a specific team member and date
def get_specific_entry(team_member_id, entry_date):
    conn = sqlite3.connect('timesheet_data.db')
    c = conn.cursor()
    
    c.execute("""
    SELECT id, hours, notes FROM timesheet_entries 
    WHERE team_member_id = ? AND entry_date = ?
    """, (team_member_id, entry_date))
    
    result = c.fetchone()
    conn.close()
    
    if result:
        return {"id": result[0], "hours": result[1], "notes": result[2]}
    return None

# Function to export data to Excel
def export_to_excel(df):
    output = BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df.to_excel(writer, index=False, sheet_name='Timesheet')
        
        # Get the worksheet and set column widths
        workbook = writer.book
        worksheet = writer.sheets['Timesheet']
        
        # Add some formatting
        header_format = workbook.add_format({
            'bold': True,
            'text_wrap': True,
            'valign': 'top',
            'bg_color': '#D9E1F2',
            'border': 1
        })
        
        # Write the column headers with the defined format
        for col_num, value in enumerate(df.columns.values):
            worksheet.write(0, col_num, value, header_format)
            worksheet.set_column(col_num, col_num, 15)
            
    output.seek(0)
    return output

# Function to generate Word document report
def generate_word_report(df, start_date, end_date):
    doc = Document()
    doc.add_heading('Team Timesheet Report', 0)
    doc.add_paragraph(f'Period: {start_date} to {end_date}')
    doc.add_paragraph('')
    
    # Summary by team member
    doc.add_heading('Summary by Team Member', level=1)
    team_summary = df.groupby('team_member')['hours'].sum().reset_index()
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Team Member'
    hdr_cells[1].text = 'Total Hours'
    
    for _, row in team_summary.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = row['team_member']
        row_cells[1].text = f"{row['hours']:.2f}"
    
    doc.add_paragraph('')
    
    # Detailed entries
    doc.add_heading('Detailed Time Entries', level=1)
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Date'
    hdr_cells[1].text = 'Team Member'
    hdr_cells[2].text = 'Hours'
    hdr_cells[3].text = 'Notes'
    
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = row['entry_date'].strftime('%Y-%m-%d')
        row_cells[1].text = row['team_member']
        row_cells[2].text = f"{row['hours']:.2f}"
        row_cells[3].text = row['notes'] if pd.notna(row['notes']) else ""
    
    # Save to BytesIO
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# Function to update entry status
def update_entry_status(entry_id, status):
    conn = sqlite3.connect('timesheet_data.db')
    c = conn.cursor()
    c.execute("UPDATE timesheet_entries SET status = ? WHERE id = ?", (status, entry_id))
    conn.commit()
    conn.close()

# Function to get weekly stats
def get_weekly_stats(start_date, end_date):
    entries = get_timesheet_entries(start_date, end_date)
    
    if entries.empty:
        return None
    
    total_hours = entries['hours'].sum()
    team_member_hours = entries.groupby('team_member')['hours'].sum().reset_index()
    daily_hours = entries.groupby(entries['entry_date'].dt.strftime('%Y-%m-%d'))['hours'].sum().reset_index()
    
    return {
        "total_hours": total_hours,
        "team_member_hours": team_member_hours,
        "daily_hours": daily_hours
    }

# Function to create a download link
def get_download_link(buffer, filename, text):
    buffer.seek(0)
    b64 = base64.b64encode(buffer.read()).decode()
    href = f'<a href="data:application/octet-stream;base64,{b64}" download="{filename}">{text}</a>'
    return href

# Initialize the database
init_db()

# Setup automatic backup system
setup_auto_backup()

# Initialize authentication
init_auth()

# Check authentication before proceeding
check_auth()

# Apply custom CSS for modern Ant Design-like styling
st.markdown("""
<style>
    /* General styling */
    .main .block-container {
        padding-top: 2rem;
        padding-bottom: 2rem;
    }

    /* Titles and headers */
    h1, h2, h3, h4, h5, h6 {
        color: #1f1f1f !important;
        font-family: 'Inter', sans-serif;
    }

    h1 {
        font-size: 32px !important;
        font-weight: 600 !important;
    }

    h2 {
        font-size: 24px !important;
        font-weight: 500 !important;
    }

    h3 {
        font-size: 20px !important;
        font-weight: 500 !important;
    }

    /* Buttons */
    .stButton > button {
        background-color: #1890ff !important;
        color: white !important;
        border-radius: 6px !important;
        border: none !important;
        padding: 8px 16px !important;
        font-size: 14px !important;
        font-weight: 500 !important;
        transition: background-color 0.2s ease;
    }

    .stButton > button:hover {
        background-color: #40a9ff !important;
    }

    /* Input fields */
    .stTextInput > div > div > input,
    .stNumberInput > div > div > input,
    .stDateInput > div > div > input,
    .stTextArea > div > div > textarea {
        border: 1px solid #d9d9d9 !important;
        border-radius: 6px !important;
        padding: 8px 12px !important;
        font-size: 14px !important;
    }

    .stTextInput > div > div > input:focus,
    .stNumberInput > div > div > input:focus,
    .stDateInput > div > div > input:focus,
    .stTextArea > div > div > textarea:focus {
        border-color: #40a9ff !important;
        box-shadow: 0 0 0 2px rgba(24, 144, 255, 0.2) !important;
    }

    /* Tabs */
    .stTabs [data-baseweb="tab-list"] {
        gap: 8px !important;
        border-bottom: 1px solid #f0f0f0 !important;
    }

    .stTabs [data-baseweb="tab"] {
        padding: 8px 16px !important;
        font-size: 14px !important;
        font-weight: 500 !important;
        color: #595959 !important;
        background-color: transparent !important;
        border-radius: 6px 6px 0 0 !important;
        transition: all 0.2s ease;
    }

    .stTabs [aria-selected="true"] {
        color: #1890ff !important;
        border-bottom: 2px solid #1890ff !important;
    }

    .stTabs [data-baseweb="tab"]:hover {
        color: #40a9ff !important;
    }

    /* Cards (for reporting metrics) */
    .reporting-card {
        background-color: white !important;
        border: 1px solid #f0f0f0 !important;
        border-radius: 8px !important;
        padding: 16px !important;
        box-shadow: 0 2px 8px rgba(0, 0, 0, 0.05) !important;
    }

    .reporting-card h4 {
        font-size: 16px !important;
        color: #595959 !important;
        margin-bottom: 8px !important;
    }

    .summary-metric {
        font-size: 24px !important;
        font-weight: 600 !important;
        color: #1f1f1f !important;
    }

    /* Tables */
    .stDataFrame {
        border-radius: 8px !important;
        box-shadow: 0 2px 8px rgba(0, 0, 0, 0.05) !important;
    }

    .stDataFrame th {
        background-color: #fafafa !important;
        font-weight: 500 !important;
        color: #1f1f1f !important;
    }

    .stDataFrame td {
        border-bottom: 1px solid #f0f0f0 !important;
    }

    /* Status badges */
    .status-approved {
        color: #52c41a !important;
        font-weight: 500 !important;
    }

    .status-pending {
        color: #faad14 !important;
        font-weight: 500 !important;
    }

    .status-rejected {
        color: #ff4d4f !important;
        font-weight: 500 !important;
    }

    /* Custom file uploader */
    .custom-file-selector input {
        display: none !important;
    }

    /* Danger zone (reset database) */
    .stWarning {
        background-color: #fff1f0 !important;
        border: 1px solid #ffccc7 !important;
        border-radius: 8px !important;
        padding: 16px !important;
    }

    .stWarning p {
        color: #ff4d4f !important;
    }
</style>
""", unsafe_allow_html=True)

# App title and description
st.title("⏱️ Team Timesheet Tracker")
st.markdown("Track team member hours easily")

# Create tabs for different app sections
tabs = st.tabs(["📝 Time Entry", "📊 Reports & Summaries", "👥 Team Management", "⚙️ Settings"])

with tabs[0]:  # Time Entry Tab
    st.header("Time Entry")
    
    # Get team members
    team_members = get_team_members()
    
    # If there are no team members, show a message
    if team_members.empty:
        st.warning("No team members available. Please add team members in the Team Management tab.")
    else:
        col1, col2 = st.columns([1, 1])
        
        with col1:
            # Team member selection
            selected_member = st.selectbox(
                "Select Team Member",
                options=team_members['id'].tolist(),
                format_func=lambda x: team_members.loc[team_members['id'] == x, 'name'].iloc[0],
                key="time_entry_member"
            )
        
        with col2:
            # Date selection with calendar
            selected_date = st.date_input(
                "Select Date", 
                value=datetime.now().date(),
                key="time_entry_date"
            )
        
        # Get existing entry if any
        existing_entry = get_specific_entry(selected_member, selected_date)
        
        # Hours and notes input
        hours_col, notes_col = st.columns([1, 3])
        
        with hours_col:
            hours = st.number_input(
                "Hours Worked", 
                min_value=0.0, 
                max_value=24.0, 
                value=existing_entry["hours"] if existing_entry else 0.0,
                step=0.5,
                format="%.1f"
            )
        
        with notes_col:
            notes = st.text_area(
                "Notes/Details", 
                value=existing_entry["notes"] if existing_entry else "",
                height=100
            )
        
        # Save button
        if st.button("Save Entry", key="save_time_entry"):
            if hours > 0:
                entry_id = save_timesheet_entry(
                    selected_member, 
                    selected_date.strftime('%Y-%m-%d'), 
                    hours, 
                    notes
                )
                st.success(f"Time entry saved successfully!")
            else:
                st.warning("Please enter hours worked.")
        
        # Show today's entries for the selected team member
        st.subheader(f"Today's Entries ({selected_date.strftime('%Y-%m-%d')})")
        
        today_entries = get_timesheet_entries(
            selected_date.strftime('%Y-%m-%d'),
            selected_date.strftime('%Y-%m-%d'),
            selected_member
        )
        
        if today_entries.empty:
            st.info("No entries for today.")
        else:
            # Create a more user-friendly view of the entries
            today_entries['Status'] = today_entries['status'].apply(
                lambda x: f'<span class="status-{x.lower()}">{x}</span>'
            )
            
            # Format the display columns
            display_df = today_entries[['hours', 'notes', 'Status']].copy()
            display_df.columns = ['Hours', 'Notes', 'Status']
            
            st.markdown(display_df.to_html(escape=False, index=False), unsafe_allow_html=True)
            
            # Calculate total hours
            total_hours = today_entries['hours'].sum()
            st.markdown(f"**Total Hours Today:** {total_hours:.1f}")

with tabs[1]:  # Reports & Summaries Tab
    st.header("Reports & Summaries")
    
    # Check if there are any team members
    team_members = get_team_members()
    if team_members.empty:
        st.warning("No team members available. Please add team members in the Team Management tab.")
    else:
        report_type = st.radio(
            "Select Report Type",
            options=["Daily View", "Weekly Summary", "Custom Date Range"],
            horizontal=True
        )
        
        if report_type == "Daily View":
            report_date = st.date_input(
                "Select Date",
                value=datetime.now().date(),
                key="report_daily_date"
            )
            start_date = end_date = report_date.strftime('%Y-%m-%d')
            
        elif report_type == "Weekly Summary":
            report_date = st.date_input(
                "Select Any Date in the Week",
                value=datetime.now().date(),
                key="report_weekly_date"
            )
            
            # Calculate the start of the week (Monday)
            start_of_week = report_date - timedelta(days=report_date.weekday())
            end_of_week = start_of_week + timedelta(days=6)
            
            start_date = start_of_week.strftime('%Y-%m-%d')
            end_date = end_of_week.strftime('%Y-%m-%d')
            
            st.info(f"Showing data for week: {start_date} to {end_date}")
            
        else:  # Custom Date Range
            col1, col2 = st.columns(2)
            with col1:
                custom_start = st.date_input(
                    "Start Date",
                    value=datetime.now().date() - timedelta(days=7),
                    key="report_custom_start"
                )
            with col2:
                custom_end = st.date_input(
                    "End Date",
                    value=datetime.now().date(),
                    key="report_custom_end"
                )
            
            start_date = custom_start.strftime('%Y-%m-%d')
            end_date = custom_end.strftime('%Y-%m-%d')
        
        # Filter for a specific team member (optional)
        include_filter = st.checkbox("Filter by Team Member", key="report_filter_member")
        
        if include_filter:
            filter_member = st.selectbox(
                "Select Team Member",
                options=team_members['id'].tolist(),
                format_func=lambda x: team_members.loc[team_members['id'] == x, 'name'].iloc[0],
                key="report_member_filter"
            )
        else:
            filter_member = None
        
        # Get entries for the selected date range
        entries = get_timesheet_entries(start_date, end_date, filter_member)
        
        if entries.empty:
            st.warning("No timesheet entries found for the selected date range.")
        else:
            # Summary stats
            total_hours = entries['hours'].sum()
            num_entries = len(entries)
            
            # Display summary metrics in cards
            metric_cols = st.columns(2)
            with metric_cols[0]:
                st.markdown(f"""
                <div class="reporting-card">
                    <h4>Total Hours</h4>
                    <div class="summary-metric">{total_hours:.1f}</div>
                </div>
                """, unsafe_allow_html=True)
                
            with metric_cols[1]:
                st.markdown(f"""
                <div class="reporting-card">
                    <h4>Team Members</h4>
                    <div class="summary-metric">{entries['team_member'].nunique()}</div>
                </div>
                """, unsafe_allow_html=True)
            
            # Create visualizations
            st.subheader("Hours by Team Member")
            team_hours = entries.groupby('team_member')['hours'].sum().reset_index()
            fig1 = px.bar(
                team_hours, 
                x='team_member', 
                y='hours',
                color='team_member',
                labels={'team_member': 'Team Member', 'hours': 'Hours'},
                height=400
            )
            st.plotly_chart(fig1, use_container_width=True)
            
            # Hours by day
            if report_type != "Daily View":
                st.subheader("Hours by Day")
                daily_hours = entries.groupby(entries['entry_date'].dt.strftime('%Y-%m-%d'))['hours'].sum().reset_index()
                daily_hours.columns = ['date', 'hours']
                fig3 = px.line(
                    daily_hours, 
                    x='date', 
                    y='hours',
                    markers=True,
                    labels={'date': 'Date', 'hours': 'Total Hours'},
                    height=400
                )
                st.plotly_chart(fig3, use_container_width=True)
            
            # Detailed entries
            st.subheader("Detailed Entries")
            
            # Format date for display
            entries['formatted_date'] = entries['entry_date'].dt.strftime('%Y-%m-%d')
            
            # Display entries in a table
            st.dataframe(
                entries[['formatted_date', 'team_member', 'hours', 'notes', 'status']].rename(
                    columns={
                        'formatted_date': 'Date',
                        'team_member': 'Team Member',
                        'hours': 'Hours',
                        'notes': 'Notes',
                        'status': 'Status'
                    }
                ),
                height=400
            )
            
            # Export options
            export_cols = st.columns(2)
            
            with export_cols[0]:
                excel_data = export_to_excel(entries[['formatted_date', 'team_member', 'hours', 'notes', 'status']])
                st.markdown(
                    get_download_link(excel_data, f"timesheet_{start_date}_to_{end_date}.xlsx", "📥 Download Excel Report"),
                    unsafe_allow_html=True
                )
                
            with export_cols[1]:
                word_doc = generate_word_report(entries, start_date, end_date)
                st.markdown(
                    get_download_link(word_doc, f"timesheet_report_{start_date}_to_{end_date}.docx", "📄 Download Word Report"),
                    unsafe_allow_html=True
                )

with tabs[2]:  # Team Management Tab
    st.header("Team Management")
    
    st.subheader("Team Members")
    
    # Force refresh team members data
    team_members = get_team_members()
    
    if team_members.empty:
        st.info("No team members found. Add a team member below.")
    else:
        st.dataframe(team_members.rename(columns={'id': 'ID', 'name': 'Name'}))
        
        # Delete team member section (only show if there are team members)
        st.subheader("Delete Team Member")
        st.warning("⚠️ Deleting a team member will also delete all their timesheet entries. This action cannot be undone.")
        
        # Select team member to delete
        delete_member_id = st.selectbox(
            "Select Team Member to Delete",
            options=team_members['id'].tolist(),
            format_func=lambda x: team_members.loc[team_members['id'] == x, 'name'].iloc[0],
            key="delete_team_member"
        )
        
        # Confirmation for deletion with the team member's name
        delete_member_name = team_members.loc[team_members['id'] == delete_member_id, 'name'].iloc[0]
        confirm_delete = st.checkbox(f"I confirm I want to delete {delete_member_name} and all their data", key="confirm_delete")
        
        if st.button("Delete Team Member", key="btn_delete_member"):
            if confirm_delete:
                conn = sqlite3.connect('timesheet_data.db')
                c = conn.cursor()
                
                try:
                    # First, delete all timesheet entries for this team member
                    c.execute("DELETE FROM timesheet_entries WHERE team_member_id = ?", (delete_member_id,))
                    
                    # Then delete the team member
                    c.execute("DELETE FROM team_members WHERE id = ?", (delete_member_id,))
                    conn.commit()
                    
                    st.success(f"Team member '{delete_member_name}' and all their data have been deleted.")
                    time.sleep(1)  # Brief pause to show the success message
                    st.rerun()  # Refresh the page to update the list
                except sqlite3.Error as e:
                    st.error(f"Error deleting team member: {e}")
                finally:
                    conn.close()
            else:
                st.error("You must confirm deletion by checking the box above.")
    
    # Add new team member
    st.subheader("Add New Team Member")
    new_member = st.text_input("Name", key="new_team_member")
    
    if st.button("Add Team Member"):
        if new_member:
            conn = sqlite3.connect('timesheet_data.db')
            c = conn.cursor()
            try:
                c.execute("INSERT INTO team_members (name) VALUES (?)", (new_member,))
                conn.commit()
                st.success(f"Team member '{new_member}' added successfully!")
                time.sleep(1)  # Brief pause to show the success message
                st.rerun()  # Refresh to show the new member
            except sqlite3.IntegrityError:
                st.error(f"Team member '{new_member}' already exists.")
            finally:
                conn.close()
        else:
            st.warning("Please enter a name.")
    
    # Entry Approval
    st.subheader("Entry Approval")
    
    # Date range selection for approvals
    col1, col2 = st.columns(2)
    with col1:
        approval_start = st.date_input(
            "Start Date",
            value=datetime.now().date() - timedelta(days=7),
            key="approval_start_date"
        )
    with col2:
        approval_end = st.date_input(
            "End Date",
            value=datetime.now().date(),
            key="approval_end_date"
        )
    
    # Get entries for approval
    approval_entries = get_timesheet_entries(
        approval_start.strftime('%Y-%m-%d'),
        approval_end.strftime('%Y-%m-%d')
    )
    
    if approval_entries.empty:
        st.warning("No timesheet entries found for the selected date range.")
    else:
        # Group by team member and date for easier management
        st.write(f"Found {len(approval_entries)} entries for review")
        
        # Allow filtering by status
        status_filter = st.multiselect(
            "Filter by Status",
            options=["Pending", "Approved", "Rejected"],
            default=["Pending"]
        )
        
        filtered_entries = approval_entries[approval_entries['status'].isin(status_filter)]
        
        if filtered_entries.empty:
            st.info(f"No entries with status: {', '.join(status_filter)}")
        else:
            # Display entries grouped by team member
            for team_member, group in filtered_entries.groupby('team_member'):
                with st.expander(f"{team_member} ({len(group)} entries)"):
                    # Display entries in a table
                    for _, entry in group.iterrows():
                        cols = st.columns([2, 1, 2])
                        with cols[0]:
                            st.write(f"**Date:** {entry['entry_date'].strftime('%Y-%m-%d')}")
                        with cols[1]:
                            st.write(f"**Hours:** {entry['hours']}")
                        with cols[2]:
                            status_options = ["Pending", "Approved", "Rejected"]
                            new_status = st.selectbox(
                                "Status",
                                options=status_options,
                                index=status_options.index(entry['status']),
                                key=f"status_{entry['id']}"
                            )
                            
                            if new_status != entry['status']:
                                if st.button("Update", key=f"update_{entry['id']}"):
                                    update_entry_status(entry['id'], new_status)
                                    st.success(f"Status updated to {new_status}")
                                    st.rerun()
                        
                        st.write(f"**Notes:** {entry['notes'] if pd.notna(entry['notes']) else ''}")
                        st.markdown("---")

with tabs[3]:  # Settings Tab
    st.header("Settings")

    # Change Password Section
    st.subheader("Change Password")
    change_password_ui()

    # Database Backup & Restore Section
    st.subheader("Database Backup & Restore")

    backup_tabs = st.tabs(["Backup", "Restore", "Backup History"])

    with backup_tabs[0]:  # Backup tab
        st.info("Create a backup of your timesheet database or download all backups.")
        
        col1, col2 = st.columns(2)
        
        with col1:
            if st.button("Create Manual Backup", key="create_manual_backup"):
                backup_path = create_manual_backup()
                st.success(f"Backup created successfully at: {backup_path}")
        
        with col2:
            if st.button("Download All Backups as ZIP", key="download_all_backups"):
                backup_zip = create_backup_zip()
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                st.markdown(
                    get_download_link(backup_zip, f"timesheet_all_backups_{timestamp}.zip", "📥 Download All Backups"),
                    unsafe_allow_html=True
                )

    with backup_tabs[1]:  # Restore tab
        st.warning("⚠️ Restoring from a backup will replace all current data with the backup's data.")
        st.info("Before restoring, a backup of the current state will be created automatically.")
        
        # Get available backups
        backups = get_available_backups()
        
        if backups.empty:
            st.warning("No backups available for restore.")
        else:
            # Display available backups
            st.write("Available Backups:")
            
            # Format the backup date for display
            backups['display_date'] = pd.to_datetime(backups['backup_date']).dt.strftime('%Y-%m-%d %H:%M:%S')
            
            # Create a user-friendly selection format
            backups['selection_label'] = backups.apply(
                lambda x: f"{x['display_date']} ({x['backup_type_display']})", 
                axis=1
            )
            
            # Create a mapping of selection label to backup path
            backup_mapping = dict(zip(backups['selection_label'], backups['backup_path']))
            
            # Let user select a backup
            selected_backup_label = st.selectbox(
                "Select a backup to restore",
                options=list(backup_mapping.keys()),
                key="restore_backup_select"
            )
            
            selected_backup_path = backup_mapping[selected_backup_label]
            
            # Confirmation
            if st.checkbox("I understand this will replace all current data with the selected backup's data.", key="confirm_restore"):
                if st.button("Restore Selected Backup", key="btn_restore_backup"):
                    result = restore_from_backup(selected_backup_path)
                    
                    if result is True:
                        st.success("Database restored successfully! Reloading page...")
                        time.sleep(2)
                        st.rerun()
                    else:
                        st.error(f"Error restoring database: {result}")

    with backup_tabs[2]:  # Backup History tab
        backups = get_available_backups()
        
        if backups.empty:
            st.info("No backup history available.")
        else:
            # Format the data for display
            display_backups = backups[['backup_date', 'backup_type_display']].copy()
            display_backups.columns = ['Backup Date', 'Type']
            display_backups['Backup Date'] = pd.to_datetime(display_backups['Backup Date']).dt.strftime('%Y-%m-%d %H:%M:%S')
            
            st.write(f"Total Backups: {len(backups)}")
            st.dataframe(display_backups)

    # Reset Database Section (Danger Zone)
    st.subheader("Reset Database")
    st.warning("⚠️ This will delete all data and reset the database to its initial state. Use with caution!")
    
    if st.checkbox("I understand that this will delete all data and cannot be undone.", key="confirm_reset_db"):
        if st.button("Reset Database", key="btn_reset_db"):
            # Create a backup before resetting
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            pre_reset_backup = f"backups/pre_reset_backup_{timestamp}.db"
            
            # Create backups directory if it doesn't exist
            if not os.path.exists('backups'):
                os.makedirs('backups')
            
            # Copy current database to backup
            shutil.copy2('timesheet_data.db', pre_reset_backup)
            
            conn = sqlite3.connect('timesheet_data.db')
            c = conn.cursor()
            
            # Log the backup
            c.execute('''
            CREATE TABLE IF NOT EXISTS backup_logs (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                backup_date TIMESTAMP,
                backup_path TEXT,
                backup_type TEXT
            )
            ''')
            
            c.execute("""
            INSERT INTO backup_logs (backup_date, backup_path, backup_type) 
            VALUES (?, ?, ?)
            """, (datetime.now().strftime('%Y-%m-%d %H:%M:%S'), pre_reset_backup, 'pre-reset'))
            
            # Drop all tables
            c.execute("DROP TABLE IF EXISTS team_members")
            c.execute("DROP TABLE IF EXISTS timesheet_entries")
            c.execute("DROP TABLE IF EXISTS auth_settings")
            c.execute("DROP TABLE IF EXISTS app_settings")
            
            # Keep the backup_logs table
            
            conn.commit()
            conn.close()
            
            # Reinitialize the database
            init_db()
            init_auth()
            
            st.success("Database has been reset to its initial state. A backup was created before reset.")
            time.sleep(2)  # Brief pause to show the success message
            st.rerun()  # Refresh the page to reflect changes
